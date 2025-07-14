#!/usr/bin/env python3
"""Generate test resumes for bulk upload testing."""

import random
import os
import zipfile
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_LEFT, TA_CENTER

# Data for generating resumes
FIRST_NAMES = [
    "John", "Jane", "Michael", "Sarah", "David", "Emily", "Robert", "Lisa",
    "James", "Mary", "William", "Jennifer", "Richard", "Linda", "Joseph", "Patricia",
    "Thomas", "Barbara", "Charles", "Elizabeth", "Christopher", "Susan", "Daniel", "Jessica",
    "Matthew", "Karen", "Anthony", "Nancy", "Donald", "Betty", "Mark", "Helen",
    "Paul", "Sandra", "Steven", "Donna", "Andrew", "Carol", "Kenneth", "Ruth",
    "Joshua", "Sharon", "Kevin", "Michelle", "Brian", "Laura", "George", "Sarah",
    "Raj", "Priya", "Ahmed", "Fatima"
]

LAST_NAMES = [
    "Smith", "Johnson", "Williams", "Brown", "Jones", "Garcia", "Miller", "Davis",
    "Rodriguez", "Martinez", "Hernandez", "Lopez", "Gonzalez", "Wilson", "Anderson", "Thomas",
    "Taylor", "Moore", "Jackson", "Martin", "Lee", "Perez", "Thompson", "White",
    "Harris", "Sanchez", "Clark", "Ramirez", "Lewis", "Robinson", "Walker", "Young",
    "Allen", "King", "Wright", "Scott", "Torres", "Nguyen", "Hill", "Flores",
    "Green", "Adams", "Nelson", "Baker", "Hall", "Rivera", "Campbell", "Mitchell",
    "Patel", "Sharma", "Chen", "Wang"
]

CITIES = [
    "New York, NY", "Los Angeles, CA", "Chicago, IL", "Houston, TX", "Phoenix, AZ",
    "Philadelphia, PA", "San Antonio, TX", "San Diego, CA", "Dallas, TX", "San Jose, CA",
    "Austin, TX", "Jacksonville, FL", "Fort Worth, TX", "Columbus, OH", "San Francisco, CA",
    "Charlotte, NC", "Indianapolis, IN", "Seattle, WA", "Denver, CO", "Washington, DC",
    "Boston, MA", "Nashville, TN", "Detroit, MI", "Portland, OR", "Las Vegas, NV",
    "Remote", "Atlanta, GA", "Miami, FL", "Oakland, CA", "Minneapolis, MN"
]

COMPANIES = [
    "Tech Corp", "Innovation Labs", "Digital Solutions Inc", "Cloud Systems", "Data Dynamics",
    "Software House", "Code Factory", "App Builders", "Web Works", "System Solutions",
    "Cyber Security Corp", "AI Innovations", "Machine Learning Co", "Big Data Inc", "Analytics Pro",
    "Mobile Apps Inc", "Enterprise Software", "Startup Hub", "Tech Startup", "Global Tech",
    "Future Systems", "Next Gen Tech", "Smart Solutions", "Tech Pioneers", "Digital Ventures"
]

UNIVERSITIES = [
    "Massachusetts Institute of Technology", "Stanford University", "Harvard University",
    "California Institute of Technology", "University of Chicago", "Princeton University",
    "Cornell University", "Yale University", "Columbia University", "University of Pennsylvania",
    "University of Michigan", "Johns Hopkins University", "Northwestern University",
    "University of California, Berkeley", "New York University", "University of California, Los Angeles",
    "Duke University", "Carnegie Mellon University", "University of Washington", "Georgia Institute of Technology"
]

JOB_POSITIONS = {
    "Senior Software Engineer": {
        "titles": ["Software Engineer", "Senior Software Engineer", "Lead Developer", "Full Stack Developer"],
        "skills": ["Python", "Java", "JavaScript", "React", "Node.js", "AWS", "Docker", "Kubernetes", "Git", "CI/CD", "Microservices", "REST APIs"],
        "experience_range": (3, 12)
    },
    "Product Manager": {
        "titles": ["Product Manager", "Senior Product Manager", "Associate Product Manager", "Product Owner"],
        "skills": ["Product Strategy", "Roadmap Planning", "Agile", "Scrum", "JIRA", "Analytics", "User Research", "A/B Testing", "Stakeholder Management"],
        "experience_range": (2, 10)
    },
    "Data Scientist": {
        "titles": ["Data Scientist", "Senior Data Scientist", "Machine Learning Engineer", "Data Analyst"],
        "skills": ["Python", "R", "Machine Learning", "Deep Learning", "TensorFlow", "PyTorch", "SQL", "Pandas", "NumPy", "Scikit-learn", "Statistics"],
        "experience_range": (1, 8)
    },
    "Frontend Developer": {
        "titles": ["Frontend Developer", "UI Developer", "React Developer", "Senior Frontend Engineer"],
        "skills": ["React", "Vue.js", "Angular", "TypeScript", "JavaScript", "HTML5", "CSS3", "Webpack", "Redux", "GraphQL", "Responsive Design"],
        "experience_range": (1, 8)
    },
    "Backend Developer": {
        "titles": ["Backend Developer", "API Developer", "Senior Backend Engineer", "Python Developer"],
        "skills": ["Python", "Java", "Node.js", "PostgreSQL", "MongoDB", "Redis", "Docker", "REST APIs", "GraphQL", "Microservices", "AWS"],
        "experience_range": (2, 10)
    },
    "DevOps Engineer": {
        "titles": ["DevOps Engineer", "Site Reliability Engineer", "Cloud Engineer", "Infrastructure Engineer"],
        "skills": ["AWS", "Azure", "GCP", "Docker", "Kubernetes", "Terraform", "CI/CD", "Jenkins", "Ansible", "Monitoring", "Linux"],
        "experience_range": (2, 10)
    }
}

DEGREES = [
    "Bachelor of Science in Computer Science",
    "Bachelor of Engineering in Software Engineering", 
    "Master of Science in Computer Science",
    "Bachelor of Science in Information Technology",
    "Master of Business Administration",
    "Bachelor of Science in Data Science",
    "Master of Science in Machine Learning",
    "Bachelor of Science in Mathematics",
    "Bachelor of Science in Physics",
    "Master of Science in Artificial Intelligence"
]

def generate_email(first_name, last_name):
    """Generate a professional email address."""
    domains = ["email.com", "mail.com", "techmail.com", "promail.com", "workmail.com"]
    formats = [
        f"{first_name.lower()}.{last_name.lower()}",
        f"{first_name.lower()}{last_name.lower()}",
        f"{first_name[0].lower()}{last_name.lower()}",
        f"{first_name.lower()}{last_name[0].lower()}"
    ]
    return f"{random.choice(formats)}@{random.choice(domains)}"

def generate_phone():
    """Generate a random phone number."""
    return f"+1-{random.randint(200, 999)}-{random.randint(200, 999)}-{random.randint(1000, 9999)}"

def generate_summary(job_position, years_exp):
    """Generate a professional summary."""
    summaries = {
        "Senior Software Engineer": [
            f"Experienced Software Engineer with {years_exp} years of expertise in full-stack development, cloud architecture, and team leadership.",
            f"Senior Software Engineer with {years_exp} years of experience building scalable applications and leading development teams.",
            f"Full-stack developer with {years_exp} years of experience in modern web technologies and distributed systems."
        ],
        "Product Manager": [
            f"Strategic Product Manager with {years_exp} years of experience driving product vision and delivering customer-centric solutions.",
            f"Results-driven Product Manager with {years_exp} years of experience in agile development and stakeholder management.",
            f"Innovative Product Manager with {years_exp} years of experience launching successful products in competitive markets."
        ],
        "Data Scientist": [
            f"Data Scientist with {years_exp} years of experience in machine learning, statistical analysis, and predictive modeling.",
            f"Analytical Data Scientist with {years_exp} years of experience transforming data into actionable insights.",
            f"Machine Learning specialist with {years_exp} years of experience developing and deploying AI solutions."
        ],
        "Frontend Developer": [
            f"Creative Frontend Developer with {years_exp} years of experience building responsive and intuitive user interfaces.",
            f"Frontend Engineer with {years_exp} years of expertise in modern JavaScript frameworks and UI/UX best practices.",
            f"UI Developer with {years_exp} years of experience creating engaging web applications and optimizing performance."
        ],
        "Backend Developer": [
            f"Backend Developer with {years_exp} years of experience designing and implementing scalable server-side applications.",
            f"Senior Backend Engineer with {years_exp} years of expertise in API development and database optimization.",
            f"Full-stack developer with {years_exp} years of backend focus, specializing in microservices architecture."
        ],
        "DevOps Engineer": [
            f"DevOps Engineer with {years_exp} years of experience automating deployments and managing cloud infrastructure.",
            f"Site Reliability Engineer with {years_exp} years of expertise in CI/CD pipelines and container orchestration.",
            f"Cloud Engineer with {years_exp} years of experience optimizing infrastructure and ensuring high availability."
        ]
    }
    return random.choice(summaries.get(job_position, [f"Professional with {years_exp} years of experience."]))

def generate_experience(job_position, years_exp):
    """Generate work experience entries."""
    job_data = JOB_POSITIONS[job_position]
    experiences = []
    remaining_years = years_exp
    
    for i in range(random.randint(2, 4)):
        if remaining_years <= 0:
            break
            
        duration = min(random.randint(1, 4), remaining_years)
        end_year = 2024 - sum(exp["duration"] for exp in experiences)
        start_year = end_year - duration
        
        title = random.choice(job_data["titles"])
        company = random.choice(COMPANIES)
        
        descriptions = {
            "Senior Software Engineer": [
                "Led development of microservices architecture serving 1M+ daily users",
                "Implemented CI/CD pipelines reducing deployment time by 70%",
                "Mentored team of 5 junior developers and conducted code reviews",
                "Designed and developed RESTful APIs for core platform features"
            ],
            "Product Manager": [
                "Managed product roadmap for flagship SaaS platform with $10M ARR",
                "Conducted user research and A/B testing to improve conversion by 25%",
                "Collaborated with engineering and design teams in agile environment",
                "Defined and tracked KPIs to measure product success"
            ],
            "Data Scientist": [
                "Developed machine learning models improving prediction accuracy by 30%",
                "Built data pipelines processing 100GB+ of daily data",
                "Created dashboards and visualizations for executive decision making",
                "Implemented recommendation system increasing user engagement by 40%"
            ],
            "Frontend Developer": [
                "Developed responsive React applications with 99.9% uptime",
                "Improved page load speed by 50% through optimization techniques",
                "Implemented design system ensuring consistent UI across products",
                "Built reusable component library used across multiple projects"
            ],
            "Backend Developer": [
                "Designed scalable backend systems handling 10K+ requests per second",
                "Optimized database queries reducing response time by 60%",
                "Implemented secure authentication and authorization systems",
                "Built event-driven architecture using message queues"
            ],
            "DevOps Engineer": [
                "Automated infrastructure provisioning using Terraform and Ansible",
                "Reduced deployment failures by 80% through improved CI/CD pipelines",
                "Implemented monitoring and alerting systems for 99.99% uptime",
                "Managed Kubernetes clusters serving production workloads"
            ]
        }
        
        experience = {
            "title": title,
            "company": company,
            "duration": duration,
            "start_year": start_year,
            "end_year": end_year,
            "description": random.sample(descriptions.get(job_position, ["Performed job duties"]), 2)
        }
        
        experiences.append(experience)
        remaining_years -= duration
    
    return experiences

def generate_education():
    """Generate education entry."""
    degree = random.choice(DEGREES)
    university = random.choice(UNIVERSITIES)
    grad_year = random.randint(2010, 2022)
    
    return {
        "degree": degree,
        "university": university,
        "year": grad_year
    }

def create_txt_resume(resume_data, filename):
    """Create a plain text resume."""
    content = []
    
    # Header
    content.append(f"{resume_data['first_name']} {resume_data['last_name']}")
    content.append(resume_data['title'])
    content.append(f"{resume_data['email']} | {resume_data['phone']} | {resume_data['location']}")
    content.append("")
    
    # Summary
    content.append("PROFESSIONAL SUMMARY")
    content.append("-" * 50)
    content.append(resume_data['summary'])
    content.append("")
    
    # Experience
    content.append("WORK EXPERIENCE")
    content.append("-" * 50)
    for exp in resume_data['experience']:
        content.append(f"{exp['title']} at {exp['company']}")
        content.append(f"{exp['start_year']} - {exp['end_year']}")
        for desc in exp['description']:
            content.append(f"• {desc}")
        content.append("")
    
    # Education
    content.append("EDUCATION")
    content.append("-" * 50)
    content.append(resume_data['education']['degree'])
    content.append(f"{resume_data['education']['university']}, {resume_data['education']['year']}")
    content.append("")
    
    # Skills
    content.append("SKILLS")
    content.append("-" * 50)
    content.append(", ".join(resume_data['skills']))
    
    with open(filename, 'w') as f:
        f.write('\n'.join(content))

def create_docx_resume(resume_data, filename):
    """Create a DOCX resume."""
    doc = Document()
    
    # Header
    header = doc.add_paragraph()
    header.add_run(f"{resume_data['first_name']} {resume_data['last_name']}\n").bold = True
    header.add_run(f"{resume_data['title']}\n")
    header.add_run(f"{resume_data['email']} | {resume_data['phone']} | {resume_data['location']}")
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Summary
    doc.add_heading('Professional Summary', level=2)
    doc.add_paragraph(resume_data['summary'])
    
    # Experience
    doc.add_heading('Work Experience', level=2)
    for exp in resume_data['experience']:
        p = doc.add_paragraph()
        p.add_run(f"{exp['title']} at {exp['company']}").bold = True
        doc.add_paragraph(f"{exp['start_year']} - {exp['end_year']}")
        for desc in exp['description']:
            doc.add_paragraph(f"• {desc}", style='List Bullet')
        doc.add_paragraph()
    
    # Education
    doc.add_heading('Education', level=2)
    doc.add_paragraph(resume_data['education']['degree'])
    doc.add_paragraph(f"{resume_data['education']['university']}, {resume_data['education']['year']}")
    
    # Skills
    doc.add_heading('Skills', level=2)
    doc.add_paragraph(", ".join(resume_data['skills']))
    
    doc.save(filename)

def create_pdf_resume(resume_data, filename):
    """Create a PDF resume."""
    doc = SimpleDocTemplate(filename, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        textColor='black',
        spaceAfter=6,
        alignment=TA_CENTER
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=12,
        textColor='black',
        spaceAfter=6
    )
    
    # Header
    story.append(Paragraph(f"{resume_data['first_name']} {resume_data['last_name']}", title_style))
    story.append(Paragraph(f"{resume_data['title']}", styles['Normal']))
    story.append(Paragraph(f"{resume_data['email']} | {resume_data['phone']} | {resume_data['location']}", styles['Normal']))
    story.append(Spacer(1, 0.2*inch))
    
    # Summary
    story.append(Paragraph("Professional Summary", heading_style))
    story.append(Paragraph(resume_data['summary'], styles['Normal']))
    story.append(Spacer(1, 0.2*inch))
    
    # Experience
    story.append(Paragraph("Work Experience", heading_style))
    for exp in resume_data['experience']:
        story.append(Paragraph(f"<b>{exp['title']} at {exp['company']}</b>", styles['Normal']))
        story.append(Paragraph(f"{exp['start_year']} - {exp['end_year']}", styles['Normal']))
        for desc in exp['description']:
            story.append(Paragraph(f"• {desc}", styles['Normal']))
        story.append(Spacer(1, 0.1*inch))
    
    # Education
    story.append(Paragraph("Education", heading_style))
    story.append(Paragraph(resume_data['education']['degree'], styles['Normal']))
    story.append(Paragraph(f"{resume_data['education']['university']}, {resume_data['education']['year']}", styles['Normal']))
    story.append(Spacer(1, 0.2*inch))
    
    # Skills
    story.append(Paragraph("Skills", heading_style))
    story.append(Paragraph(", ".join(resume_data['skills']), styles['Normal']))
    
    doc.build(story)

def generate_resume_data(job_position):
    """Generate resume data for a specific job position."""
    first_name = random.choice(FIRST_NAMES)
    last_name = random.choice(LAST_NAMES)
    
    job_data = JOB_POSITIONS[job_position]
    years_exp = random.randint(*job_data["experience_range"])
    
    resume_data = {
        "first_name": first_name,
        "last_name": last_name,
        "email": generate_email(first_name, last_name),
        "phone": generate_phone(),
        "location": random.choice(CITIES),
        "title": random.choice(job_data["titles"]),
        "summary": generate_summary(job_position, years_exp),
        "experience": generate_experience(job_position, years_exp),
        "education": generate_education(),
        "skills": random.sample(job_data["skills"], min(8, len(job_data["skills"]))),
        "years_experience": years_exp,
        "job_position": job_position
    }
    
    return resume_data

def main():
    """Generate test resumes."""
    # Define how many resumes per position
    position_counts = {
        "Senior Software Engineer": 10,
        "Product Manager": 8,
        "Data Scientist": 8,
        "Frontend Developer": 8,
        "Backend Developer": 8,
        "DevOps Engineer": 8
    }
    
    # Create directories
    individual_dir = "resumes/individual"
    bulk_dir = "resumes/bulk"
    
    all_resumes = []
    position_resumes = {pos: [] for pos in position_counts.keys()}
    
    # Generate resumes
    resume_count = 0
    for position, count in position_counts.items():
        for i in range(count):
            resume_count += 1
            resume_data = generate_resume_data(position)
            
            # Determine format
            if resume_count <= 20:
                format_type = "pdf"
            elif resume_count <= 40:
                format_type = "docx"
            else:
                format_type = "txt"
            
            # Generate filename
            filename = f"{resume_data['first_name']}_{resume_data['last_name']}_{position.replace(' ', '_')}_{resume_count}.{format_type}"
            filepath = os.path.join(individual_dir, filename)
            
            # Create resume file
            if format_type == "txt":
                create_txt_resume(resume_data, filepath)
            elif format_type == "docx":
                create_docx_resume(resume_data, filepath)
            elif format_type == "pdf":
                create_pdf_resume(resume_data, filepath)
            
            all_resumes.append(filepath)
            position_resumes[position].append(filepath)
            
            print(f"Created {filepath}")
    
    # Create ZIP files
    # All resumes ZIP
    with zipfile.ZipFile(os.path.join(bulk_dir, "all_resumes.zip"), 'w') as zipf:
        for resume_path in all_resumes:
            zipf.write(resume_path, os.path.basename(resume_path))
    print(f"Created {bulk_dir}/all_resumes.zip with {len(all_resumes)} resumes")
    
    # Position-specific ZIPs
    for position, resumes in position_resumes.items():
        if resumes:
            zip_name = f"{position.lower().replace(' ', '_')}.zip"
            with zipfile.ZipFile(os.path.join(bulk_dir, zip_name), 'w') as zipf:
                for resume_path in resumes:
                    zipf.write(resume_path, os.path.basename(resume_path))
            print(f"Created {bulk_dir}/{zip_name} with {len(resumes)} resumes")
    
    # Create mixed small batches
    for i in range(5):
        batch_resumes = random.sample(all_resumes, min(10, len(all_resumes)))
        with zipfile.ZipFile(os.path.join(bulk_dir, f"batch_{i+1}.zip"), 'w') as zipf:
            for resume_path in batch_resumes:
                zipf.write(resume_path, os.path.basename(resume_path))
        print(f"Created {bulk_dir}/batch_{i+1}.zip with {len(batch_resumes)} resumes")

if __name__ == "__main__":
    main()