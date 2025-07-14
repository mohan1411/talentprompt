"""File parsing service for extracting text from various file formats."""

import io
import re
from typing import Optional, Tuple

import PyPDF2
from docx import Document


class FileParser:
    """Service for parsing different file types."""
    
    ALLOWED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.txt'}
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
    
    @classmethod
    def validate_file(cls, filename: str, file_size: int) -> Tuple[bool, Optional[str]]:
        """Validate file before processing."""
        # Check file extension
        ext = filename.lower().split('.')[-1] if '.' in filename else ''
        if f'.{ext}' not in cls.ALLOWED_EXTENSIONS:
            return False, f"File type .{ext} not allowed. Allowed types: {', '.join(cls.ALLOWED_EXTENSIONS)}"
        
        # Check file size
        if file_size > cls.MAX_FILE_SIZE:
            return False, f"File size exceeds maximum allowed size of {cls.MAX_FILE_SIZE / 1024 / 1024}MB"
        
        return True, None
    
    @staticmethod
    def extract_text_from_pdf(file_content: bytes) -> str:
        """Extract text from PDF file."""
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text_parts = []
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            
            return '\n'.join(text_parts)
        except Exception as e:
            raise ValueError(f"Failed to parse PDF: {str(e)}")
    
    @staticmethod
    def extract_text_from_docx(file_content: bytes) -> str:
        """Extract text from DOCX file."""
        try:
            doc = Document(io.BytesIO(file_content))
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            return '\n'.join(text_parts)
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX: {str(e)}")
    
    @staticmethod
    def extract_text_from_txt(file_content: bytes) -> str:
        """Extract text from TXT file."""
        try:
            return file_content.decode('utf-8', errors='ignore')
        except Exception as e:
            raise ValueError(f"Failed to parse TXT: {str(e)}")
    
    @classmethod
    def extract_text(cls, file_content: bytes, filename: str) -> str:
        """Extract text from file based on its extension."""
        ext = filename.lower().split('.')[-1] if '.' in filename else ''
        
        if ext == 'pdf':
            text = cls.extract_text_from_pdf(file_content)
        elif ext in ['docx', 'doc']:
            text = cls.extract_text_from_docx(file_content)
        elif ext == 'txt':
            text = cls.extract_text_from_txt(file_content)
        else:
            raise ValueError(f"Unsupported file type: .{ext}")
        
        # Clean up text
        text = re.sub(r'\s+', ' ', text)  # Replace multiple spaces with single space
        text = text.strip()
        
        if not text:
            raise ValueError("No text content found in file")
        
        return text